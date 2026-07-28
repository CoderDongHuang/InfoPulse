"""Portable report exporters. Optional format dependencies are loaded lazily."""

import html
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2] / "storage" / "exports"


def markdown_lines(markdown: str):
    for raw in markdown.splitlines():
        match = re.match(r"^(#{1,3})\s+(.+)", raw)
        yield (len(match.group(1)), match.group(2)) if match else (0, raw)


def reference_lines(citations: list[dict]) -> list[str]:
    return [f"[{index}] {item['title']} - {item['url']}" for index, item in enumerate(citations, 1)]


def export_report(report, version, citations: list[dict], fmt: str) -> Path:
    ROOT.mkdir(parents=True, exist_ok=True)
    suffix = "md" if fmt == "markdown" else fmt
    path = ROOT / f"{report.id}-{version.id}.{suffix}"
    references = reference_lines(citations)
    content = version.content_markdown.rstrip() + "\n\n## 来源引用\n" + "\n".join(references)

    if fmt == "markdown":
        path.write_text(content, encoding="utf-8")
    elif fmt == "html":
        _export_html(path, report.title, content)
    elif fmt == "docx":
        _export_docx(path, report, version, citations)
    elif fmt == "pdf":
        _export_pdf(path, report, version, citations)
    else:
        raise ValueError(f"Unsupported export format: {fmt}")
    return path


def _export_html(path: Path, title: str, content: str) -> None:
    blocks = []
    for level, text in markdown_lines(content):
        escaped = html.escape(text)
        blocks.append(f"<h{level}>{escaped}</h{level}>" if level else (f"<p>{escaped}</p>" if text else ""))
    document = f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width"><title>{html.escape(title)}</title>
<style>body{{max-width:820px;margin:48px auto;padding:0 24px;color:#17231f;font:16px/1.75 Arial,'Microsoft YaHei',sans-serif}}h1,h2,h3{{color:#173f59}}h1{{border-bottom:3px solid #188178;padding-bottom:16px}}h2{{margin-top:32px}}p{{white-space:pre-wrap}}@media print{{body{{margin:0}}}}</style></head>
<body><h1>{html.escape(title)}</h1>{''.join(blocks)}</body></html>"""
    path.write_text(document, encoding="utf-8")


def _export_docx(path: Path, report, version, citations: list[dict]) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("DOCX export requires python-docx") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin, section.bottom_margin = Inches(0.8), Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.85)
    normal = document.styles["Normal"]
    normal.font.name, normal.font.size = "Microsoft YaHei", Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.3
    for name, size in (("Title", 26), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 11)):
        style = document.styles[name]
        style.font.name, style.font.size, style.font.color.rgb = "Microsoft YaHei", Pt(size), RGBColor(23, 63, 89)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.text, header.alignment = "InfoPulse · AI Intelligence Report", WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size, header.runs[0].font.color.rgb = Pt(8), RGBColor(100, 116, 110)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("InfoPulse  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    document.add_heading(report.title, 0)
    document.add_paragraph(f"{report.report_type.upper()}  ·  VERSION {version.version_number}")
    document.add_heading("目录", 1)
    for level, text in markdown_lines(version.content_markdown):
        if level:
            document.add_paragraph(f"{'  ' * max(level - 1, 0)}{text}")
    document.add_page_break()
    for level, text in markdown_lines(version.content_markdown):
        if level:
            document.add_heading(text, level=min(level, 3))
        elif text:
            document.add_paragraph(text)
    _add_docx_charts(document, version.structured_content or {})
    document.add_heading("来源引用", 1)
    for line in reference_lines(citations):
        document.add_paragraph(line)
    document.save(path)


def _add_docx_charts(document, structured: dict) -> None:
    for chart in structured.get("charts", []):
        document.add_heading(str(chart.get("title", "数据图表")), 2)
        labels, values = chart.get("labels", []), chart.get("values", [])
        if labels and values:
            table = document.add_table(rows=1, cols=2)
            table.style = "Light Shading Accent 1"
            table.rows[0].cells[0].text, table.rows[0].cells[1].text = "项目", "数值"
            for label, value in zip(labels, values):
                cells = table.add_row().cells
                cells[0].text, cells[1].text = str(label), str(value)


def _export_pdf(path: Path, report, version, citations: list[dict]) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        from reportlab.lib import colors
    except ImportError as exc:
        raise RuntimeError("PDF export requires reportlab") from exc

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    base = getSampleStyleSheet()
    body = ParagraphStyle("CN", parent=base["BodyText"], fontName="STSong-Light", fontSize=10, leading=16, spaceAfter=7)
    heading = {n: ParagraphStyle(f"CNH{n}", parent=body, fontSize=24 - n * 4, leading=28 - n * 3, textColor="#173F59", spaceBefore=12, spaceAfter=8) for n in (1, 2, 3)}
    story = [Paragraph(html.escape(report.title), heading[1]), Paragraph(f"{report.report_type.upper()} · VERSION {version.version_number}", body), Spacer(1, 8), Paragraph("目录", heading[2])]
    story.extend(Paragraph(html.escape(text), body) for level, text in markdown_lines(version.content_markdown) if level)
    story.append(PageBreak())
    for level, text in markdown_lines(version.content_markdown):
        if text:
            story.append(Paragraph(html.escape(text), heading[min(level, 3)] if level else body))
    for chart in (version.structured_content or {}).get("charts", []):
        story.append(Paragraph(html.escape(str(chart.get("title", "数据图表"))), heading[2]))
        rows = [["项目", "数值"], *zip(chart.get("labels", []), chart.get("values", []))]
        table = Table(rows, colWidths=[90 * mm, 60 * mm])
        table.setStyle(TableStyle([("FONTNAME", (0, 0), (-1, -1), "STSong-Light"), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DDECEA")), ("GRID", (0, 0), (-1, -1), .25, colors.HexColor("#AABBB5")), ("PADDING", (0, 0), (-1, -1), 6)]))
        story.append(table)
    story.append(Paragraph("来源引用", heading[2]))
    story.extend(Paragraph(html.escape(line), body) for line in reference_lines(citations))

    def decorate(canvas, doc):
        canvas.saveState()
        canvas.setFont("STSong-Light", 8)
        canvas.setFillColor("#66746E")
        canvas.drawString(20 * mm, 12 * mm, "InfoPulse · AI Intelligence Report")
        canvas.drawRightString(190 * mm, 12 * mm, str(doc.page))
        canvas.restoreState()

    SimpleDocTemplate(str(path), pagesize=A4, rightMargin=20 * mm, leftMargin=20 * mm, topMargin=20 * mm, bottomMargin=20 * mm).build(story, onFirstPage=decorate, onLaterPages=decorate)
